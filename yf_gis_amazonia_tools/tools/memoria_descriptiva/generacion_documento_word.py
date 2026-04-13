# -*- coding: utf-8 -*-
"""
Generación del documento Word - Memoria Descriptiva v3.0
Formato profesional justificado, tablas con bandas de color.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

_MESES = {'January':'enero','February':'febrero','March':'marzo','April':'abril',
          'May':'mayo','June':'junio','July':'julio','August':'agosto',
          'September':'setiembre','October':'octubre','November':'noviembre','December':'diciembre'}

def _fecha_es():
    s = datetime.now().strftime('%d de %B de %Y')
    for en, es in _MESES.items(): s = s.replace(en, es)
    return s

def _fnum(v, fmt):
    try: return format(float(v), fmt)
    except: return str(v)

# ── XML helpers ───────────────────────────────────────────────────────────────

def _cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def _cell_borders(cell, color='AAAAAA', size=4):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement('w:{}'.format(side))
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(size))
        b.set(qn('w:space'),'0');    b.set(qn('w:color'),color)
        tcB.append(b)
    tcPr.append(tcB)

def _para_border_bottom(p, color='17375E', size='6'):
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),size)
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)

# ── Párrafo justificado ───────────────────────────────────────────────────────

def _pj(doc, text, size=11, bold=False, spb=0, spa=6, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(spb); pf.space_after = Pt(spa)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent: pf.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.name = 'Arial'
    return p

# ── Encabezado de sección ─────────────────────────────────────────────────────

def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Arial'; r.font.bold = True
    if level == 1:
        r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x17,0x37,0x5E)
        _para_border_bottom(p, '17375E', '6')
    else:
        r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x2E,0x6E,0x3E)
        _para_border_bottom(p, '2E6E3E', '4')
    return p

# ── Tabla helpers ─────────────────────────────────────────────────────────────

def _hdr_row(row, bg='17375E', fg='FFFFFF'):
    for cell in row.cells:
        _cell_bg(cell, bg); _cell_borders(cell, bg, 4)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            for r in para.runs:
                r.font.bold = True; r.font.name = 'Arial'; r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(int(fg[0:2],16),int(fg[2:4],16),int(fg[4:6],16))

def _data_row(row, alt=False, center=False):
    bg = 'EAF4EA' if alt else 'FFFFFF'
    for cell in row.cells:
        _cell_bg(cell, bg); _cell_borders(cell, 'CCCCCC', 2)
        for para in cell.paragraphs:
            if center: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            for r in para.runs:
                r.font.name = 'Arial'; r.font.size = Pt(9)

# =============================================================================
# FUNCIÓN PRINCIPAL
# =============================================================================

def generar_documento_word(datos_formulario, datos_procesados, sufijo_archivo=None):
    """
    Genera un documento Word de Memoria Descriptiva.

    Args:
        datos_formulario : dict con datos del formulario
        datos_procesados : dict con vertices, area, perimetro, colindantes, etc.
        sufijo_archivo   : str para distinguir archivos en modo atlas

    Returns:
        Ruta del archivo generado
    """
    doc = Document()

    for section in doc.sections:
        section.page_width    = Cm(21.0); section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.5);  section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0);  section.right_margin  = Cm(2.5)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    _encabezado(doc, datos_formulario, datos_procesados)
    _s_solicitante(doc, datos_formulario)
    _s_generalidades(doc, datos_formulario.get('generalidades',''))
    _s_ubicacion(doc, datos_formulario.get('ubicacion',{}))
    _s_colindantes(doc, datos_procesados.get('colindantes',{}))
    _s_tecnica(doc, datos_procesados)
    _s_mapa(doc, datos_formulario.get('info_mapa',{}))
    _s_firma(doc, datos_formulario)

    # ── Ruta de salida ────────────────────────────────────────────────────────
    base_path = datos_formulario['output_file']
    if sufijo_archivo:
        base, ext = os.path.splitext(base_path)
        ext = ext.lstrip('.') or 'docx'
        clean = "".join(c for c in sufijo_archivo
                        if c.isalnum() or c in (' ','_','-')).strip().replace(' ','_')
        output_path = "{}_{}.{}".format(base, clean[:50], ext)
    else:
        output_path = base_path

    # ── Evitar sobreescritura: agregar _2, _3, etc. si el archivo ya existe ──
    if os.path.exists(output_path):
        base2, ext2 = os.path.splitext(output_path)
        counter = 2
        while os.path.exists("{}_{}{}".format(base2, counter, ext2)):
            counter += 1
        output_path = "{}_{}{}".format(base2, counter, ext2)
        print("  Nombre duplicado - guardando como: {}".format(os.path.basename(output_path)))

    doc.save(output_path)
    return output_path


# =============================================================================
# SECCIONES
# =============================================================================

def _encabezado(doc, df, dp):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('MEMORIA DESCRIPTIVA')
    r.font.name='Arial'; r.font.size=Pt(16); r.font.bold=True
    r.font.color.rgb = RGBColor(0x17,0x37,0x5E)

    # Subtítulo con nombre del predio/propietario
    nombre_predio = dp.get('nombre_propietario','') or df.get('ubicacion',{}).get('sector','')
    if nombre_predio:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run('PREDIO DE: {}'.format(nombre_predio.upper()))
        r2.font.name='Arial'; r2.font.size=Pt(12); r2.font.bold=True
        r2.font.color.rgb = RGBColor(0x2E,0x6E,0x3E)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before=Pt(4); p3.paragraph_format.space_after=Pt(10)
    _para_border_bottom(p3,'17375E','12')


def _s_solicitante(doc, df):
    # Tomar el nombre del propietario del predio procesado si existe en datos_procesados
    nombre = df.get('_nombre_propietario_actual','') or df.get('solicitante',{}).get('nombre','')
    dni    = df.get('_dni_actual','')                or df.get('solicitante',{}).get('dni','')

    _heading(doc, 'I.   DATOS DEL SOLICITANTE', 1)
    for label, val in [
        ('Nombre y Apellidos : ', (nombre or 'No especificado').upper()),
        ('D.N.I.                      : ', dni or 'No especificado'),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        p.paragraph_format.left_indent=Cm(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1=p.add_run(label); r1.font.name='Arial'; r1.font.size=Pt(11)
        r2=p.add_run(val);   r2.font.name='Arial'; r2.font.size=Pt(11); r2.font.bold=True


def _s_generalidades(doc, texto):
    _heading(doc, 'II.  GENERALIDADES', 1)
    if not texto or not texto.strip():
        texto = ("La presente Memoria Descriptiva tiene por finalidad describir las "
                 "características técnicas del predio materia del presente trámite, "
                 "determinando sus linderos, medidas perimétricas, área total, colindantes "
                 "y demás aspectos técnicos que permitan su correcta identificación y "
                 "ubicación en el territorio nacional, de acuerdo con las normas "
                 "técnicas vigentes del SERFOR.")
    for i, par in enumerate([p.strip() for p in texto.split('\n') if p.strip()]):
        _pj(doc, par, spb=4 if i==0 else 2, spa=4)


def _s_ubicacion(doc, dat):
    _heading(doc, 'III. UBICACIÓN', 1)
    sector=dat.get('sector',''); distrito=dat.get('distrito','')
    provincia=dat.get('provincia',''); depto=dat.get('departamento','')
    zona=dat.get('zona','')

    partes = []
    if sector:    partes.append('el sector denominado {}'.format(sector))
    if distrito:  partes.append('el Distrito de {}'.format(distrito))
    if provincia: partes.append('la Provincia de {}'.format(provincia))
    if depto:     partes.append('el Departamento de {}'.format(depto))

    txt = ('El predio se encuentra ubicado en {}, República del Perú.'.format(', '.join(partes))
           if partes else 'El predio se encuentra ubicado en la República del Perú.')
    _pj(doc, txt, spb=4, spa=6)

    items = [(k,v) for k,v in [
        ('Sector / Localidad',sector),('Zona UTM',zona),('Distrito',distrito),
        ('Provincia',provincia),('Departamento',depto),('País','Perú')] if v]
    if items:
        t = doc.add_table(rows=1, cols=2)
        t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
        hdr=t.rows[0]; hdr.cells[0].text='DESCRIPCIÓN'; hdr.cells[1].text='DETALLE'
        _hdr_row(hdr)
        for i,(k,v) in enumerate(items):
            row=t.add_row(); row.cells[0].text=k; row.cells[1].text=str(v)
            _data_row(row, i%2==1)
            for c in row.cells:
                c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
                for r in c.paragraphs[0].runs: r.font.name='Arial'; r.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)


def _s_colindantes(doc, dat):
    _heading(doc, 'IV.  COLINDANTES', 1)
    orden=['NORTE','SUR','ESTE','OESTE']
    partes=[]
    for lado in orden:
        info=dat.get(lado,{}); nombre=info.get('nombre','Terrenos del Estado') if isinstance(info,dict) else str(info)
        partes.append('por el {} con {}'.format(lado.capitalize(), nombre))
    if partes:
        txt=('El predio colinda: '+'; '.join(partes[:-1])+(' y '+partes[-1] if len(partes)>1 else partes[0])+'.')
        _pj(doc, txt, spb=4, spa=6)

    t=doc.add_table(rows=1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    hdr=t.rows[0]; hdr.cells[0].text='LADO'; hdr.cells[1].text='COLINDANTE'; hdr.cells[2].text='OBSERVACIÓN'
    _hdr_row(hdr)
    for i,lado in enumerate(orden):
        info=dat.get(lado,{}); nombre=info.get('nombre','Terrenos del Estado') if isinstance(info,dict) else str(info)
        obs=info.get('observacion','') if isinstance(info,dict) else ''
        row=t.add_row(); row.cells[0].text='POR EL {}'.format(lado)
        row.cells[1].text=str(nombre); row.cells[2].text=str(obs)
        _data_row(row, i%2==1)
        row.cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        try: row.cells[0].paragraphs[0].runs[0].font.bold=True
        except: pass
        for c in row.cells:
            for r in c.paragraphs[0].runs: r.font.name='Arial'; r.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)


def _s_tecnica(doc, dp):
    _heading(doc, 'V.   INFORMACIÓN TÉCNICA DEL PREDIO', 1)
    vertices=dp.get('vertices',[]); area=dp.get('area',0); perimetro=dp.get('perimetro',0)
    desc=dp.get('descripcion_linderos',''); fuente_area=dp.get('fuente_area','')

    # 5.1
    _heading(doc, '5.1.   Linderos y Medidas Perimétricas', 2)
    _pj(doc,'El predio se enmarca con las siguientes medidas perimétricas, expresadas '
        'en el Sistema UTM (Universal Transversal de Mercator), con coordenadas en metros:',
        spb=4, spa=4)
    if desc: _pj(doc, desc, spb=2, spa=6)

    # 5.2
    _heading(doc, '5.2.   Cuadro Técnico de Vértices', 2)
    _pj(doc,'A continuación, se presenta el cuadro técnico con las coordenadas de los '
        'vértices del predio, junto con las distancias y azimuts de cada lado:',
        spb=4, spa=6)

    if vertices:
        t=doc.add_table(rows=1,cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
        hdr=t.rows[0]
        for i,h in enumerate(['VÉRTICE','LADO','ESTE (m)','NORTE (m)','DISTANCIA (m)','AZIMUT (°)']):
            hdr.cells[i].text=h
        _hdr_row(hdr)
        for idx,v in enumerate(vertices):
            row=t.add_row()
            row.cells[0].text=str(v.get('vertice','V{:02d}'.format(idx+1)))
            row.cells[1].text=str(v.get('lado',''))
            row.cells[2].text=_fnum(v.get('este',0),',.4f')
            row.cells[3].text=_fnum(v.get('norte',0),',.4f')
            row.cells[4].text=_fnum(v.get('distancia',0),'.2f')
            row.cells[5].text=_fnum(v.get('azimut',0),'.4f')
            _data_row(row, idx%2==1, center=True)
    else:
        _pj(doc,'No se encontraron vértices. Verifique el campo de relación entre polígonos y puntos.',spb=4,spa=6)

    doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # 5.3
    _heading(doc, '5.3.   Área y Perímetro', 2)
    try:
        ah=float(area); am2=round(ah*10000,2)
        nota_fuente=' [Fuente: {}]'.format(fuente_area) if fuente_area else ''
        txt=('El predio tiene una superficie total de {:,.4f} hectáreas '
             '({:,.2f} m²) y un perímetro de {:,.2f} metros lineales.{}'.format(
             ah, am2, float(perimetro), nota_fuente))
    except: txt='Los datos de área y perímetro no están disponibles.'
    _pj(doc, txt, spb=4, spa=6)

    t2=doc.add_table(rows=2,cols=2); t2.alignment=WD_TABLE_ALIGNMENT.CENTER; t2.style='Table Grid'
    hdr2=t2.rows[0]; hdr2.cells[0].text='ÁREA TOTAL'; hdr2.cells[1].text='PERÍMETRO TOTAL'
    _hdr_row(hdr2,'2E6E3E')
    dr=t2.rows[1]
    dr.cells[0].text='{} ha'.format(_fnum(area,',.4f'))
    dr.cells[1].text='{} m'.format(_fnum(perimetro,',.2f'))
    _cell_bg(dr.cells[0],'EAF4EA'); _cell_bg(dr.cells[1],'EAF4EA')
    for c in dr.cells:
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in c.paragraphs[0].runs:
            r.font.name='Arial'; r.font.size=Pt(12); r.font.bold=True
        _cell_borders(c,'2E6E3E',4)
    doc.add_paragraph().paragraph_format.space_after=Pt(8)


def _s_mapa(doc, info):
    _heading(doc, 'VI.  INFORMACIÓN TÉCNICA DEL MAPA', 1)
    _pj(doc,'El presente plano ha sido elaborado utilizando el sistema de referencia '
        'geodésico y la proyección cartográfica indicados a continuación:',spb=4,spa=6)
    if info:
        t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
        hdr=t.rows[0]; hdr.cells[0].text='PARÁMETRO'; hdr.cells[1].text='VALOR'
        _hdr_row(hdr)
        for i,(d,v) in enumerate(info.items()):
            if v:
                row=t.add_row(); row.cells[0].text=str(d); row.cells[1].text=str(v)
                _data_row(row,i%2==1)
                for c in row.cells:
                    c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
                    for r in c.paragraphs[0].runs: r.font.name='Arial'; r.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after=Pt(8)


def _s_firma(doc, df):
    p_sep=doc.add_paragraph(); p_sep.paragraph_format.space_before=Pt(18)
    _para_border_bottom(p_sep,'17375E','4')

    p_f=doc.add_paragraph(); p_f.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p_f.paragraph_format.space_before=Pt(8); p_f.paragraph_format.space_after=Pt(28)
    r=p_f.add_run('Puerto Maldonado, {}'.format(_fecha_es()))
    r.font.name='Arial'; r.font.size=Pt(10); r.font.italic=True

    for _ in range(4):
        sp=doc.add_paragraph(''); sp.paragraph_format.line_spacing=Pt(13)

    # Nombre y DNI pueden venir de datos_procesados (propietario real) o del formulario
    nombre = (df.get('_nombre_propietario_actual','') or
              df.get('solicitante',{}).get('nombre','') or 'SOLICITANTE')
    dni    = df.get('_dni_actual','') or df.get('solicitante',{}).get('dni','')

    for txt,bold,size,color in [
        ('_'*40, False, 11, None),
        (nombre.upper(), True, 11, None),
        ('D.N.I. N\u00b0 {}'.format(dni) if dni else '', False, 10, None),
        ('PROPIETARIO / SOLICITANTE', True, 10, RGBColor(0x17,0x37,0x5E)),
    ]:
        if not txt: continue
        p=doc.add_paragraph(txt); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(2)
        for r in p.runs:
            r.font.name='Arial'; r.font.size=Pt(size); r.font.bold=bold
            if color: r.font.color.rgb=color
