# -*- coding: utf-8 -*-
"""
/***************************************************************************
 ReportGenerator
                                 A QGIS plugin
 Búsqueda avanzada de atributos con funcionalidades similares a ArcMap
                             -------------------
        begin                : 2025-04-21
        copyright            : (C) 2025 by Yuri Caller
        email                : yuricaller@gmail.com
 ***************************************************************************/

/***************************************************************************
 *                                                                         *
 *   This program is free software; you can redistribute it and/or modify  *
 *   it under the terms of the GNU General Public License as published by  *
 *   the Free Software Foundation; either version 2 of the License, or     *
 *   (at your option) any later version.                                   *
 *                                                                         *
 ***************************************************************************/
"""

import os
import tempfile
from datetime import datetime

try:
    import matplotlib.pyplot as plt
    import numpy as np
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

from qgis.PyQt.QtCore import QObject, pyqtSignal, QSettings, QSize
from qgis.core import (
    QgsProject, QgsVectorLayer, QgsFeature, QgsGeometry, QgsRectangle,
    QgsCoordinateReferenceSystem, QgsCoordinateTransform, QgsWkbTypes,
    QgsMapSettings, QgsMapRendererParallelJob, QgsLayoutExporter
)

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from docx2pdf import convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

class ReportGenerator(QObject):
    """Generator for reports from search results."""
    
    reportStarted = pyqtSignal()
    reportProgress = pyqtSignal(int, str)  # progress percentage, status message
    reportCompleted = pyqtSignal(str)  # output path
    reportError = pyqtSignal(str)  # error message
    
    def __init__(self, iface, parent=None):
        """Constructor."""
        super(ReportGenerator, self).__init__(parent)
        self.iface = iface
        self.temp_dir = tempfile.mkdtemp(prefix="qgis_attribute_search_")
        
        # Check if python-docx is installed
        if not HAS_DOCX:
            self.reportError.emit("La biblioteca python-docx no está instalada. Por favor, instálela para habilitar la generación de reportes.")
    
    def generate_report(self, results, output_path, report_params):
        """Generate report based on parameters."""
        if not HAS_DOCX:
            self.reportError.emit("La biblioteca python-docx no está instalada. Por favor, instálela para habilitar la generación de reportes.")
            return False
        
        if not results:
            self.reportError.emit("No hay resultados para generar un reporte.")
            return False
        
        # Start report generation
        self.reportStarted.emit()
        self.reportProgress.emit(0, "Iniciando generación de reporte...")
        
        try:
            # Extract report parameters
            report_title = report_params.get('title', f"Reporte de Búsqueda - {datetime.now().strftime('%d/%m/%Y')}")
            report_type = report_params.get('type', 'detailed')
            selected_fields = report_params.get('fields', [])
            include_map = report_params.get('include_map', True)
            include_attributes = report_params.get('include_attributes', True)
            include_charts = report_params.get('include_charts', True)
            include_metadata = report_params.get('include_metadata', True)
            company_logo = report_params.get('company_logo', None)
            author_name = report_params.get('author_name', "QGIS Attribute Search Plugin")
            
            # Determine output format
            _, ext = os.path.splitext(output_path)
            output_format = ext.lower().replace('.', '')
            
            # Generate Word document
            if output_format == 'docx':
                self.generate_word_report(
                    results,
                    output_path,
                    report_title,
                    report_type,
                    selected_fields,
                    include_map,
                    include_attributes,
                    include_charts,
                    include_metadata,
                    company_logo,
                    author_name
                )
                
                self.reportCompleted.emit(output_path)
                return True
            
            elif output_format == 'pdf':
                # For PDF, we first generate a Word document, then convert it
                temp_docx = os.path.join(self.temp_dir, "temp_report.docx")
                
                self.generate_word_report(
                    results,
                    temp_docx,
                    report_title,
                    report_type,
                    selected_fields,
                    include_map,
                    include_attributes,
                    include_charts,
                    include_metadata,
                    company_logo,
                    author_name
                )
                
                # Convert to PDF
                self.convert_to_pdf(temp_docx, output_path)
                
                self.reportCompleted.emit(output_path)
                return True
            
            else:
                self.reportError.emit(f"Formato de salida no soportado: {output_format}")
                return False
        
        except Exception as e:
            self.reportError.emit(f"Error al generar el reporte: {str(e)}")
            return False
    
    def generate_word_report(self, results, output_path, report_title, report_type, selected_fields,
                           include_map, include_attributes, include_charts, include_metadata,
                           company_logo, author_name):
        """Generate Word report."""
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = report_title
        doc.core_properties.author = author_name
        
        # Add company logo if provided
        if company_logo and os.path.exists(company_logo):
            self.reportProgress.emit(5, "Añadiendo logo de empresa...")
            doc.add_picture(company_logo, width=Inches(2))
        
        # Add title
        self.reportProgress.emit(10, "Añadiendo título y encabezado...")
        title = doc.add_heading(report_title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        date_run.italic = True
        
        # Add summary
        doc.add_heading("Resumen", level=1)
        summary_paragraph = doc.add_paragraph()
        summary_paragraph.add_run(f"Este reporte contiene los resultados de una búsqueda de atributos en QGIS. ")
        
        # Count unique features and layers
        unique_features = set((r['layer'].id(), r['feature_id']) for r in results)
        unique_layers = set(r['layer'] for r in results)
        
        summary_paragraph.add_run(f"Se encontraron {len(unique_features)} entidades en {len(unique_layers)} capas.")
        
        # Add map if requested
        if include_map:
            self.reportProgress.emit(20, "Generando mapa de ubicación...")
            doc.add_heading("Mapa de Ubicación", level=1)
            
            # Generate map image
            map_image_path = self.generate_map_image(results)
            
            if map_image_path:
                doc.add_picture(map_image_path, width=Inches(6))
                
                # Add caption
                map_caption = doc.add_paragraph("Mapa de ubicación de las entidades encontradas.")
                map_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                map_caption.style = 'Caption'
        
        # Add attributes if requested
        if include_attributes:
            self.reportProgress.emit(40, "Añadiendo información de atributos...")
            
            if report_type == "detailed":
                self.add_detailed_attributes(doc, results, selected_fields)
            elif report_type == "summary":
                self.add_summary_attributes(doc, results, selected_fields)
            else:
                self.add_single_entity_attributes(doc, results, selected_fields)
        
        # Add charts if requested
        if include_charts and report_type != "single":
            self.reportProgress.emit(60, "Generando gráficos estadísticos...")
            doc.add_heading("Gráficos Estadísticos", level=1)
            
            # Generate charts
            chart_paths = self.generate_charts(results, selected_fields)
            
            for field_name, chart_path in chart_paths.items():
                if chart_path:
                    doc.add_heading(f"Distribución de valores para {field_name}", level=2)
                    doc.add_picture(chart_path, width=Inches(6))
        
        # Add metadata if requested
        if include_metadata:
            self.reportProgress.emit(80, "Añadiendo metadatos de las capas...")
            doc.add_heading("Metadatos de las Capas", level=1)
            
            # Get unique layers
            layers = set(r['layer'] for r in results)
            
            for layer in layers:
                doc.add_heading(layer.name(), level=2)
                
                metadata_paragraph = doc.add_paragraph()
                
                # Add basic metadata
                metadata_paragraph.add_run(f"Tipo de geometría: ").bold = True
                metadata_paragraph.add_run(self.get_geometry_type_name(layer.geometryType()))
                metadata_paragraph.add_run("\n")
                
                metadata_paragraph.add_run(f"Sistema de coordenadas: ").bold = True
                metadata_paragraph.add_run(layer.crs().authid())
                metadata_paragraph.add_run("\n")
                
                metadata_paragraph.add_run(f"Número de entidades: ").bold = True
                metadata_paragraph.add_run(str(layer.featureCount()))
                metadata_paragraph.add_run("\n")
                
                metadata_paragraph.add_run(f"Campos: ").bold = True
                field_names = [field.name() for field in layer.fields()]
                metadata_paragraph.add_run(", ".join(field_names))
        
        # Save document
        self.reportProgress.emit(90, "Guardando documento...")
        doc.save(output_path)
        
        self.reportProgress.emit(100, "Reporte generado exitosamente.")
    
    def add_detailed_attributes(self, doc, results, selected_fields):
        """Add detailed attributes table to the document."""
        doc.add_heading("Resultados Detallados", level=1)
        
        # Group results by layer
        results_by_layer = {}
        for result in results:
            layer = result['layer']
            if layer not in results_by_layer:
                results_by_layer[layer] = []
            results_by_layer[layer].append(result)
        
        # Add tables for each layer
        for layer, layer_results in results_by_layer.items():
            doc.add_heading(f"Capa: {layer.name()}", level=2)
            
            # Create table
            # Header row + one row per feature
            unique_features = set(r['feature_id'] for r in layer_results)
            table = doc.add_table(rows=1, cols=len(selected_fields) + 1)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "ID"
            for i, field_name in enumerate(selected_fields):
                header_cells[i + 1].text = field_name
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add data rows
            for feature_id in unique_features:
                # Get feature
                request = QgsFeatureRequest().setFilterFid(feature_id)
                features = list(layer.getFeatures(request))
                
                if not features:
                    continue
                
                feature = features[0]
                
                # Add row
                row_cells = table.add_row().cells
                row_cells[0].text = str(feature_id)
                
                # Add field values
                for i, field_name in enumerate(selected_fields):
                    if field_name in feature.fields().names():
                        row_cells[i + 1].text = str(feature[field_name])
                    else:
                        row_cells[i + 1].text = ""
            
            # Add space after table
            doc.add_paragraph()
    
    def add_summary_attributes(self, doc, results, selected_fields):
        """Add summary attributes to the document."""
        doc.add_heading("Resumen de Resultados", level=1)
        
        # Group results by layer
        results_by_layer = {}
        for result in results:
            layer = result['layer']
            if layer not in results_by_layer:
                results_by_layer[layer] = []
            results_by_layer[layer].append(result)
        
        # Add summary for each layer
        for layer, layer_results in results_by_layer.items():
            doc.add_heading(f"Capa: {layer.name()}", level=2)
            
            # Count unique features
            unique_features = set(r['feature_id'] for r in layer_results)
            
            summary_paragraph = doc.add_paragraph()
            summary_paragraph.add_run(f"Número de entidades encontradas: ").bold = True
            summary_paragraph.add_run(str(len(unique_features)))
            summary_paragraph.add_run("\n")
            
            # Add field statistics
            for field_name in selected_fields:
                if field_name not in layer.fields().names():
                    continue
                
                doc.add_heading(f"Estadísticas para el campo {field_name}", level=3)
                
                # Get field values
                field_values = []
                for feature_id in unique_features:
                    request = QgsFeatureRequest().setFilterFid(feature_id)
                    features = list(layer.getFeatures(request))
                    
                    if features:
                        feature = features[0]
                        field_values.append(str(feature[field_name]))
                
                # Count unique values
                value_counts = {}
                for value in field_values:
                    if value in value_counts:
                        value_counts[value] += 1
                    else:
                        value_counts[value] = 1
                
                # Create table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Valor"
                header_cells[1].text = "Conteo"
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for value, count in sorted(value_counts.items(), key=lambda x: x[1], reverse=True):
                    row_cells = table.add_row().cells
                    row_cells[0].text = value
                    row_cells[1].text = str(count)
                
                # Add space after table
                doc.add_paragraph()
                
                # Try to calculate numeric statistics
                numeric_values = []
                for value in field_values:
                    try:
                        numeric_values.append(float(value))
                    except (ValueError, TypeError):
                        pass
                
                if numeric_values:
                    stats_paragraph = doc.add_paragraph()
                    stats_paragraph.add_run("Estadísticas numéricas:\n").bold = True
                    stats_paragraph.add_run(f"Mínimo: {min(numeric_values)}\n")
                    stats_paragraph.add_run(f"Máximo: {max(numeric_values)}\n")
                    stats_paragraph.add_run(f"Promedio: {sum(numeric_values) / len(numeric_values):.2f}\n")
                    stats_paragraph.add_run(f"Mediana: {sorted(numeric_values)[len(numeric_values) // 2]}\n")
                    if len(numeric_values) > 1:
                        stats_paragraph.add_run(f"Desviación estándar: {np.std(numeric_values):.2f}")
    
    def add_single_entity_attributes(self, doc, results, selected_fields):
        """Add single entity attributes to the document."""
        doc.add_heading("Ficha de Entidad", level=1)
        
        # Get first result
        if not results:
            return
        
        result = results[0]
        layer = result['layer']
        feature_id = result['feature_id']
        
        # Get feature
        request = QgsFeatureRequest().setFilterFid(feature_id)
        features = list(layer.getFeatures(request))
        
        if not features:
            return
        
        feature = features[0]
        
        # Add layer info
        doc.add_heading(f"Capa: {layer.name()}", level=2)
        
        # Add feature ID
        id_paragraph = doc.add_paragraph()
        id_paragraph.add_run(f"ID de Entidad: ").bold = True
        id_paragraph.add_run(str(feature_id))
        
        # Add attributes
        doc.add_heading("Atributos", level=2)
        
        # Create table
        table = doc.add_table(rows=len(selected_fields), cols=2)
        table.style = 'Table Grid'
        
        # Add rows
        for i, field_name in enumerate(selected_fields):
            if field_name in feature.fields().names():
                row_cells = table.rows[i].cells
                row_cells[0].text = field_name
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].text = str(feature[field_name])
        
        # Add feature geometry info
        doc.add_heading("Geometría", level=2)
        
        geometry_paragraph = doc.add_paragraph()
        
        geometry = feature.geometry()
        if geometry:
            geometry_paragraph.add_run(f"Tipo: ").bold = True
            geometry_paragraph.add_run(self.get_geometry_type_name(layer.geometryType()))
            geometry_paragraph.add_run("\n")
            
            if layer.geometryType() == QgsWkbTypes.PolygonGeometry:
                # Add area
                area = geometry.area()
                if area < 10000:
                    area_str = f"{area:.2f} m²"
                else:
                    area_str = f"{area/10000:.4f} ha"
                
                geometry_paragraph.add_run(f"Área: ").bold = True
                geometry_paragraph.add_run(area_str)
                geometry_paragraph.add_run("\n")
                
                # Add perimeter
                perimeter = geometry.length()
                if perimeter < 1000:
                    perimeter_str = f"{perimeter:.2f} m"
                else:
                    perimeter_str = f"{perimeter/1000:.4f} km"
                
                geometry_paragraph.add_run(f"Perímetro: ").bold = True
                geometry_paragraph.add_run(perimeter_str)
            
            elif layer.geometryType() == QgsWkbTypes.LineGeometry:
                # Add length
                length = geometry.length()
                if length < 1000:
                    length_str = f"{length:.2f} m"
                else:
                    length_str = f"{length/1000:.4f} km"
                
                geometry_paragraph.add_run(f"Longitud: ").bold = True
                geometry_paragraph.add_run(length_str)
            
            elif layer.geometryType() == QgsWkbTypes.PointGeometry:
                # Add coordinates
                point = geometry.asPoint()
                
                geometry_paragraph.add_run(f"Coordenadas: ").bold = True
                geometry_paragraph.add_run(f"X: {point.x():.6f}, Y: {point.y():.6f}")
    
    def generate_map_image(self, results):
        """Generate map image of search results."""
        # Create temporary file for map image
        map_image_path = os.path.join(self.temp_dir, "map_image.png")
        
        # Get all features from results
        features_by_layer = {}
        for result in results:
            layer = result['layer']
            feature_id = result['feature_id']
            
            if layer not in features_by_layer:
                features_by_layer[layer] = []
            
            features_by_layer[layer].append(feature_id)
        
        # Calculate combined extent
        combined_extent = None
        
        for layer, feature_ids in features_by_layer.items():
            for feature_id in feature_ids:
                # Get feature
                request = QgsFeatureRequest().setFilterFid(feature_id)
                features = list(layer.getFeatures(request))
                
                if features:
                    feature = features[0]
                    geometry = feature.geometry()
                    
                    if geometry:
                        # Get feature extent
                        extent = geometry.boundingBox()
                        
                        # Transform extent to project CRS if needed
                        if layer.crs() != QgsProject.instance().crs():
                            transform = QgsCoordinateTransform(
                                layer.crs(),
                                QgsProject.instance().crs(),
                                QgsProject.instance()
                            )
                            extent = transform.transformBoundingBox(extent)
                        
                        # Combine extents
                        if combined_extent is None:
                            combined_extent = extent
                        else:
                            combined_extent.combineExtentWith(extent)
        
        if not combined_extent:
            return None
        
        # Add buffer for better visibility
        combined_extent.grow(combined_extent.width() * 0.1)
        
        # Create map settings
        map_settings = QgsMapSettings()
        map_settings.setDestinationCrs(QgsProject.instance().crs())
        map_settings.setExtent(combined_extent)
        
        # Set output size
        map_settings.setOutputSize(QSize(800, 600))
        
        # Add all visible layers
        layers = QgsProject.instance().mapLayers().values()
        map_settings.setLayers(layers)
        
        # Create map renderer
        render_job = QgsMapRendererParallelJob(map_settings)
        render_job.start()
        render_job.waitForFinished()
        
        # Get rendered image
        image = render_job.renderedImage()
        
        # Save image
        image.save(map_image_path, "PNG")
        
        return map_image_path
    
    def generate_charts(self, results, selected_fields):
        """Generate charts for selected fields."""
        chart_paths = {}
        
        # Group results by layer
        results_by_layer = {}
        for result in results:
            layer = result['layer']
            if layer not in results_by_layer:
                results_by_layer[layer] = []
            results_by_layer[layer].append(result)
        
        # Generate charts for each field
        for field_name in selected_fields:
            # Check if field exists in any layer
            field_exists = False
            for layer in results_by_layer.keys():
                if field_name in layer.fields().names():
                    field_exists = True
                    break
            
            if not field_exists:
                continue
            
            # Count values across all layers
            value_counts = {}
            
            for layer, layer_results in results_by_layer.items():
                if field_name not in layer.fields().names():
                    continue
                
                # Get unique features
                unique_features = set(r['feature_id'] for r in layer_results)
                
                for feature_id in unique_features:
                    request = QgsFeatureRequest().setFilterFid(feature_id)
                    features = list(layer.getFeatures(request))
                    
                    if features:
                        feature = features[0]
                        value = str(feature[field_name])
                        
                        if value in value_counts:
                            value_counts[value] += 1
                        else:
                            value_counts[value] = 1
            
            # Skip if no values
            if not value_counts:
                continue
            
            # Sort by count
            sorted_items = sorted(value_counts.items(), key=lambda x: x[1], reverse=True)
            
            # Limit to top 10 for readability
            if len(sorted_items) > 10:
                sorted_items = sorted_items[:10]
                other_count = sum(value_counts[k] for k in value_counts if k not in [x[0] for x in sorted_items])
                if other_count > 0:
                    sorted_items.append(("Otros", other_count))
            
            # Extract labels and values
            labels = [item[0] for item in sorted_items]
            values = [item[1] for item in sorted_items]
            
            # Create chart
            plt.figure(figsize=(8, 6))
            
            # Create bar chart
            plt.bar(labels, values)
            plt.xlabel('Valores')
            plt.ylabel('Conteo')
            plt.title(f'Distribución de valores para el campo "{field_name}"')
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            
            # Save chart
            chart_path = os.path.join(self.temp_dir, f"chart_{field_name}.png")
            plt.savefig(chart_path)
            plt.close()
            
            chart_paths[field_name] = chart_path
        
        return chart_paths
    
    def convert_to_pdf(self, docx_path, pdf_path):
        """Convert Word document to PDF."""
        if HAS_DOCX2PDF:
            self.reportProgress.emit(95, "Convirtiendo a PDF...")
            convert(docx_path, pdf_path)
        else:
            self.reportError.emit(
                "La biblioteca docx2pdf no está instalada. "
                "El reporte se ha guardado en formato Word."
            )
            # Copy the Word file to the requested PDF path but with .docx extension
            import shutil
            pdf_base = os.path.splitext(pdf_path)[0]
            shutil.copy2(docx_path, f"{pdf_base}.docx")
    
    def get_geometry_type_name(self, geometry_type):
        """Get human-readable name for geometry type."""
        if geometry_type == QgsWkbTypes.PointGeometry:
            return "Punto"
        elif geometry_type == QgsWkbTypes.LineGeometry:
            return "Línea"
        elif geometry_type == QgsWkbTypes.PolygonGeometry:
            return "Polígono"
        else:
            return "Desconocido"
