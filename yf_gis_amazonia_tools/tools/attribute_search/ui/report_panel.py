# -*- coding: utf-8 -*-
"""
/***************************************************************************
 ReportPanel
                                 A QGIS plugin
 Búsqueda avanzada de atributos con funcionalidades similares a ArcMap
                             -------------------
        begin                : 2025-04-21
        copyright            : (C) 2025 by Yuri Caller
        email                : yuricaller@gmail.com
 ***************************************************************************/

/***************************************************************************
 *                                                                         *
 *   This program is free software; you can redistribute it and/or modify  *
 *   it under the terms of the GNU General Public License as published by  *
 *   the Free Software Foundation; either version 2 of the License, or     *
 *   (at your option) any later version.                                   *
 *                                                                         *
 ***************************************************************************/
"""

import os
import tempfile
from datetime import datetime

from qgis.PyQt.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox, QPushButton,
    QGroupBox, QCheckBox, QLineEdit, QFileDialog, QMessageBox, QProgressBar,
    QRadioButton, QListWidget, QListWidgetItem, QAbstractItemView, QSplitter,
    QFrame, QGridLayout, QSpacerItem, QSizePolicy
)
from qgis.PyQt.QtCore import Qt, QSize
from qgis.PyQt.QtGui import QIcon, QFont, QPixmap

from qgis.core import (
    QgsProject, QgsVectorLayer, QgsFeature, QgsGeometry, QgsRectangle,
    QgsCoordinateReferenceSystem, QgsCoordinateTransform, QgsWkbTypes,
    QgsMapSettings, QgsMapRendererParallelJob, QgsLayoutExporter, QgsFeatureRequest # Added QgsFeatureRequest
)
from qgis.gui import QgsMapCanvas, QgsRubberBand

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

class ReportPanel(QWidget):
    """Panel for generating reports from search results."""
    
    def __init__(self, iface):
        """Constructor."""
        super(ReportPanel, self).__init__()
        self.iface = iface
        self.plugin_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
        
        # Initialize variables
        self.results = []
        self.temp_dir = tempfile.mkdtemp(prefix="qgis_attribute_search_")
        
        # Set up the user interface
        self.setup_ui()
    
    def setup_ui(self):
        """Set up the user interface."""
        # Main layout
        self.main_layout = QVBoxLayout(self)
        self.main_layout.setContentsMargins(6, 6, 6, 6)
        
        # Check if python-docx is installed
        if not HAS_DOCX:
            self.warning_label = QLabel(
                "La biblioteca python-docx no está instalada. "
                "Por favor, instálela para habilitar la generación de reportes."
            )
            self.warning_label.setStyleSheet("color: red;")
            self.main_layout.addWidget(self.warning_label)
            
            self.install_button = QPushButton("Instalar python-docx")
            self.install_button.clicked.connect(self.install_docx)
            self.main_layout.addWidget(self.install_button)
            
            return
        
        # Report configuration
        self.config_group = QGroupBox("Configuración del Reporte")
        self.config_layout = QVBoxLayout(self.config_group)
        
        # Report title
        self.title_layout = QHBoxLayout()
        self.title_label = QLabel("Título del Reporte:")
        self.title_layout.addWidget(self.title_label)
        
        self.title_input = QLineEdit()
        self.title_input.setPlaceholderText("Reporte de Búsqueda de Atributos")
        self.title_layout.addWidget(self.title_input)
        
        self.config_layout.addLayout(self.title_layout)
        
        # Report type
        self.type_group = QGroupBox("Tipo de Reporte")
        self.type_layout = QVBoxLayout(self.type_group)
        
        self.detailed_radio = QRadioButton("Reporte Detallado (todas las entidades)")
        self.detailed_radio.setChecked(True)
        self.type_layout.addWidget(self.detailed_radio)
        
        self.summary_radio = QRadioButton("Reporte Resumido (estadísticas)")
        self.type_layout.addWidget(self.summary_radio)
        
        self.single_radio = QRadioButton("Ficha de Entidad (entidad única)")
        self.type_layout.addWidget(self.single_radio)
        
        self.config_layout.addWidget(self.type_group)
        
        # Content options
        self.content_group = QGroupBox("Contenido del Reporte")
        self.content_layout = QVBoxLayout(self.content_group)
        
        self.include_map_check = QCheckBox("Incluir mapa de ubicación")
        self.include_map_check.setChecked(True)
        self.content_layout.addWidget(self.include_map_check)
        
        self.include_attributes_check = QCheckBox("Incluir tabla de atributos")
        self.include_attributes_check.setChecked(True)
        self.content_layout.addWidget(self.include_attributes_check)
        
        self.include_charts_check = QCheckBox("Incluir gráficos estadísticos")
        self.include_charts_check.setChecked(True)
        self.content_layout.addWidget(self.include_charts_check)
        
        self.include_metadata_check = QCheckBox("Incluir metadatos de la capa")
        self.include_metadata_check.setChecked(True)
        self.content_layout.addWidget(self.include_metadata_check)
        
        self.config_layout.addWidget(self.content_group)
        
        # Field selection
        self.field_group = QGroupBox("Campos a Incluir")
        self.field_layout = QVBoxLayout(self.field_group)
        
        self.field_list = QListWidget()
        self.field_list.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.field_layout.addWidget(self.field_list)
        
        self.field_buttons_layout = QHBoxLayout()
        
        self.select_all_fields_button = QPushButton("Seleccionar Todo")
        self.select_all_fields_button.clicked.connect(self.select_all_fields)
        self.field_buttons_layout.addWidget(self.select_all_fields_button)
        
        self.clear_fields_button = QPushButton("Limpiar Selección")
        self.clear_fields_button.clicked.connect(self.clear_field_selection)
        self.field_buttons_layout.addWidget(self.clear_fields_button)
        
        self.field_layout.addLayout(self.field_buttons_layout)
        
        self.config_layout.addWidget(self.field_group)
        
        self.main_layout.addWidget(self.config_group)
        
        # Output options
        self.output_group = QGroupBox("Opciones de Salida")
        self.output_layout = QGridLayout(self.output_group)
        
        # Format selection
        self.format_label = QLabel("Formato:")
        self.output_layout.addWidget(self.format_label, 0, 0)
        
        self.format_combo = QComboBox()
        self.format_combo.addItems(["Word (.docx)", "PDF"])
        self.output_layout.addWidget(self.format_combo, 0, 1)
        
        # Output path
        self.output_path_label = QLabel("Guardar en:")
        self.output_layout.addWidget(self.output_path_label, 1, 0)
        
        self.output_path_layout = QHBoxLayout()
        
        self.output_path_input = QLineEdit()
        self.output_path_input.setPlaceholderText("Seleccione ubicación para guardar...")
        self.output_path_layout.addWidget(self.output_path_input)
        
        self.browse_button = QPushButton("Examinar...")
        self.browse_button.clicked.connect(self.browse_output_path)
        self.output_path_layout.addWidget(self.browse_button)
        
        self.output_layout.addLayout(self.output_path_layout, 1, 1)
        
        self.main_layout.addWidget(self.output_group)
        
        # Generate button
        self.generate_layout = QHBoxLayout()
        
        self.generate_button = QPushButton("Generar Reporte")
        self.generate_button.setIcon(QIcon(os.path.join(self.plugin_dir, "icon.png")))
        self.generate_button.clicked.connect(self.generate_report)
        self.generate_layout.addWidget(self.generate_button)
        
        self.main_layout.addLayout(self.generate_layout)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.main_layout.addWidget(self.progress_bar)
    
    def install_docx(self):
        """Install python-docx package."""
        try:
            import pip
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            
            # Install python-docx
            pip.main(['install', 'python-docx'])
            
            self.progress_bar.setValue(90)
            
            # Try to import again
            import docx
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            self.progress_bar.setValue(100)
            
            # Reload UI
            self.warning_label.setVisible(False)
            self.install_button.setVisible(False)
            self.setup_ui()
            
            QMessageBox.information(
                self,
                "Éxito",
                "La biblioteca python-docx se ha instalado correctamente."
            )
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Error al instalar python-docx: {str(e)}\n\n"
                "Por favor, instálela manualmente ejecutando:\n"
                "pip install python-docx"
            )
        finally:
            self.progress_bar.setVisible(False)
    
    def prepare_report(self):
        """Prepare report panel with current results."""
        # Get results from parent dialog
        parent = self.parent()
        while parent and not hasattr(parent, 'results_panel'):
            parent = parent.parent()
        
        if parent and hasattr(parent, 'results_panel'):
            self.results = parent.results_panel.results
        
        # Update field list
        self.update_field_list()
        
        # Set default title
        if not self.title_input.text():
            self.title_input.setText(f"Reporte de Búsqueda - {datetime.now().strftime('%d/%m/%Y')}")
    
    def update_field_list(self):
        """Update field list with available fields."""
        self.field_list.clear()
        
        if not self.results:
            return
        
        # Get unique field names from all layers
        field_names = set()
        for result in self.results:
            layer = result['layer']
            for field in layer.fields():
                field_names.add(field.name())
        
        # Add field names to list
        for field_name in sorted(field_names):
            item = QListWidgetItem(field_name)
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
            item.setCheckState(Qt.Checked)
            self.field_list.addItem(item)
    
    def select_all_fields(self):
        """Select all fields in the list."""
        for i in range(self.field_list.count()):
            item = self.field_list.item(i)
            item.setCheckState(Qt.Checked)
    
    def clear_field_selection(self):
        """Clear field selection."""
        for i in range(self.field_list.count()):
            item = self.field_list.item(i)
            item.setCheckState(Qt.Unchecked)
    
    def browse_output_path(self):
        """Browse for output file path."""
        # Get selected format
        format_text = self.format_combo.currentText()
        if "Word" in format_text:
            file_filter = "Word Document (*.docx)"
            default_ext = ".docx"
        else:
            file_filter = "PDF Document (*.pdf)"
            default_ext = ".pdf"
        
        # Get file path
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar Reporte",
            "",
            file_filter
        )
        
        if file_path:
            # Add extension if missing
            if not file_path.lower().endswith(default_ext):
                file_path += default_ext
            
            self.output_path_input.setText(file_path)
    
    def generate_report(self):
        """Generate report based on current settings."""
        # Check if there are results
        if not self.results:
            QMessageBox.warning(
                self,
                "Advertencia",
                "No hay resultados para generar un reporte."
            )
            return
        
        # Check if output path is specified
        output_path = self.output_path_input.text()
        if not output_path:
            QMessageBox.warning(
                self,
                "Advertencia",
                "Por favor, especifique una ubicación para guardar el reporte."
            )
            return
        
        # Get selected fields
        selected_fields = []
        for i in range(self.field_list.count()):
            item = self.field_list.item(i)
            if item.checkState() == Qt.Checked:
                selected_fields.append(item.text())
        
        if not selected_fields:
            QMessageBox.warning(
                self,
                "Advertencia",
                "Por favor, seleccione al menos un campo para incluir en el reporte."
            )
            return
        
        # Show progress bar
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        try:
            # Get report type
            if self.detailed_radio.isChecked():
                report_type = "detailed"
            elif self.summary_radio.isChecked():
                report_type = "summary"
            else:
                report_type = "single"
            
            # Get content options
            include_map = self.include_map_check.isChecked()
            include_attributes = self.include_attributes_check.isChecked()
            include_charts = self.include_charts_check.isChecked()
            include_metadata = self.include_metadata_check.isChecked()
            
            # Get report title
            report_title = self.title_input.text()
            if not report_title:
                report_title = f"Reporte de Búsqueda - {datetime.now().strftime('%d/%m/%Y')}"
            
            # Get output format
            format_text = self.format_combo.currentText()
            if "Word" in format_text:
                output_format = "docx"
            else:
                output_format = "pdf"
            
            # Generate report
            self.progress_bar.setValue(10)
            
            if output_format == "docx":
                self.generate_word_report(
                    output_path,
                    report_title,
                    report_type,
                    selected_fields,
                    include_map,
                    include_attributes,
                    include_charts,
                    include_metadata
                )
            else:
                # For PDF, we first generate a Word document, then convert it
                temp_docx = os.path.join(self.temp_dir, "temp_report.docx")
                self.generate_word_report(
                    temp_docx,
                    report_title,
                    report_type,
                    selected_fields,
                    include_map,
                    include_attributes,
                    include_charts,
                    include_metadata
                )
                
                # Convert to PDF
                self.convert_to_pdf(temp_docx, output_path)
            
            self.progress_bar.setValue(100)
            
            # Show success message
            QMessageBox.information(
                self,
                "Éxito",
                f"Reporte generado exitosamente en:\n{output_path}"
            )
            
            # Open the file
            import subprocess
            import sys
            
            if sys.platform == 'win32':
                os.startfile(output_path)
            elif sys.platform == 'darwin':
                subprocess.call(['open', output_path])
            else:
                subprocess.call(['xdg-open', output_path])
        
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Error al generar el reporte: {str(e)}"
            )
        
        finally:
            self.progress_bar.setVisible(False)
    
    def generate_word_report(self, output_path, report_title, report_type, selected_fields,
                            include_map, include_attributes, include_charts, include_metadata):
        """Generate Word report."""
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = report_title
        doc.core_properties.author = "QGIS Attribute Search Plugin"
        
        # Add title
        title = doc.add_heading(report_title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        date_run.italic = True
        
        # Add summary
        doc.add_heading("Resumen", level=1)
        summary_paragraph = doc.add_paragraph()
        summary_paragraph.add_run(f"Este reporte contiene los resultados de una búsqueda de atributos en QGIS. ")
        summary_paragraph.add_run(f"Se encontraron {len(self.results)} coincidencias en {len(set(r['layer'] for r in self.results))} capas.")
        
        # Progress update
        self.progress_bar.setValue(20)
        
        # Add map if requested
        if include_map:
            doc.add_heading("Mapa de Ubicación", level=1)
            
            # Generate map image
            map_image_path = self.generate_map_image()
            
            if map_image_path:
                doc.add_picture(map_image_path, width=Inches(6))
                
                # Add caption
                map_caption = doc.add_paragraph("Mapa de ubicación de las entidades encontradas.")
                map_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                map_caption.style = 'Caption'
        
        # Progress update
        self.progress_bar.setValue(40)
        
        # Add attributes if requested
        if include_attributes:
            if report_type == "detailed":
                self.add_detailed_attributes(doc, selected_fields)
            elif report_type == "summary":
                self.add_summary_attributes(doc, selected_fields)
            else:
                self.add_single_entity_attributes(doc, selected_fields)
        
        # Progress update
        self.progress_bar.setValue(60)
        
        # Add charts if requested
        if include_charts and report_type != "single":
            doc.add_heading("Gráficos Estadísticos", level=1)
            
            # Generate charts
            chart_paths = self.generate_charts(selected_fields)
            
            for field_name, chart_path in chart_paths.items():
                if chart_path:
                    doc.add_heading(f"Distribución de valores para {field_name}", level=2)
                    doc.add_picture(chart_path, width=Inches(6))
        
        # Progress update
        self.progress_bar.setValue(80)
        
        # Add metadata if requested
        if include_metadata:
            doc.add_heading("Metadatos de las Capas", level=1)
            
            # Get unique layers
            layers = set(r['layer'] for r in self.results)
            
            for layer in layers:
                doc.add_heading(layer.name(), level=2)
                
                metadata_paragraph = doc.add_paragraph()
                
                # Add basic metadata
                metadata_paragraph.add_run(f"Tipo de geometría: ").bold = True
                metadata_paragraph.add_run(self.get_geometry_type_name(layer.geometryType()))
                metadata_paragraph.add_run("\n")
                
                metadata_paragraph.add_run(f"Sistema de coordenadas: ").bold = True
                metadata_paragraph.add_run(layer.crs().authid())
                metadata_paragraph.add_run("\n")
                
                metadata_paragraph.add_run(f"Número de entidades: ").bold = True
                metadata_paragraph.add_run(str(layer.featureCount()))
                metadata_paragraph.add_run("\n")
                
                metadata_paragraph.add_run(f"Campos: ").bold = True
                field_names = [field.name() for field in layer.fields()]
                metadata_paragraph.add_run(", ".join(field_names))
        
        # Save document
        doc.save(output_path)
    
    def add_detailed_attributes(self, doc, selected_fields):
        """Add detailed attributes table to the document."""
        doc.add_heading("Resultados Detallados", level=1)
        
        # Group results by layer
        results_by_layer = {}
        for result in self.results:
            layer = result['layer']
            if layer not in results_by_layer:
                results_by_layer[layer] = []
            results_by_layer[layer].append(result)
        
        # Add tables for each layer
        for layer, layer_results in results_by_layer.items():
            doc.add_heading(f"Capa: {layer.name()}", level=2)
            
            # Create table
            # Header row + one row per feature
            unique_features = set(r['feature_id'] for r in layer_results)
            table = doc.add_table(rows=1, cols=len(selected_fields) + 1)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "ID"
            for i, field_name in enumerate(selected_fields):
                header_cells[i + 1].text = field_name
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add data rows
            for feature_id in unique_features:
                # Get feature
                request = QgsFeatureRequest().setFilterFid(feature_id)
                features = list(layer.getFeatures(request))
                
                if not features:
                    continue
                
                feature = features[0]
                
                # Add row
                row_cells = table.add_row().cells
                row_cells[0].text = str(feature_id)
                
                # Add field values
                for i, field_name in enumerate(selected_fields):
                    if field_name in feature.fields().names():
                        row_cells[i + 1].text = str(feature[field_name])
                    else:
                        row_cells[i + 1].text = ""
            
            # Add space after table
            doc.add_paragraph()
    
    def add_summary_attributes(self, doc, selected_fields):
        """Add summary attributes to the document."""
        doc.add_heading("Resumen de Resultados", level=1)
        
        # Group results by layer
        results_by_layer = {}
        for result in self.results:
            layer = result['layer']
            if layer not in results_by_layer:
                results_by_layer[layer] = []
            results_by_layer[layer].append(result)
        
        # Add summary for each layer
        for layer, layer_results in results_by_layer.items():
            doc.add_heading(f"Capa: {layer.name()}", level=2)
            
            # Count unique features
            unique_features = set(r['feature_id'] for r in layer_results)
            
            summary_paragraph = doc.add_paragraph()
            summary_paragraph.add_run(f"Número de entidades encontradas: ").bold = True
            summary_paragraph.add_run(str(len(unique_features)))
            summary_paragraph.add_run("\n")
            
            # Add field statistics
            for field_name in selected_fields:
                if field_name not in layer.fields().names():
                    continue
                
                doc.add_heading(f"Estadísticas para el campo {field_name}", level=3)
                
                # Get field values
                field_values = []
                for feature_id in unique_features:
                    request = QgsFeatureRequest().setFilterFid(feature_id)
                    features = list(layer.getFeatures(request))
                    
                    if features:
                        feature = features[0]
                        field_values.append(str(feature[field_name]))
                
                # Count unique values
                value_counts = {}
                for value in field_values:
                    if value in value_counts:
                        value_counts[value] += 1
                    else:
                        value_counts[value] = 1
                
                # Create table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Valor"
                header_cells[1].text = "Conteo"
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for value, count in sorted(value_counts.items(), key=lambda x: x[1], reverse=True):
                    row_cells = table.add_row().cells
                    row_cells[0].text = value
                    row_cells[1].text = str(count)
                
                # Add space after table
                doc.add_paragraph()
    
    def add_single_entity_attributes(self, doc, selected_fields):
        """Add single entity attributes to the document."""
        doc.add_heading("Ficha de Entidad", level=1)
        
        # Get first result
        if not self.results:
            return
        
        result = self.results[0]
        layer = result['layer']
        feature_id = result['feature_id']
        
        # Get feature
        request = QgsFeatureRequest().setFilterFid(feature_id)
        features = list(layer.getFeatures(request))
        
        if not features:
            return
        
        feature = features[0]
        
        # Add layer info
        doc.add_heading(f"Capa: {layer.name()}", level=2)
        
        # Add feature ID
        id_paragraph = doc.add_paragraph()
        id_paragraph.add_run(f"ID de Entidad: ").bold = True
        id_paragraph.add_run(str(feature_id))
        
        # Add attributes
        doc.add_heading("Atributos", level=2)
        
        # Create table
        table = doc.add_table(rows=len(selected_fields), cols=2)
        table.style = 'Table Grid'
        
        # Add rows
        for i, field_name in enumerate(selected_fields):
            if field_name in feature.fields().names():
                row_cells = table.rows[i].cells
                row_cells[0].text = field_name
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].text = str(feature[field_name])
        
        # Add feature geometry info
        doc.add_heading("Geometría", level=2)
        
        geometry_paragraph = doc.add_paragraph()
        
        geometry = feature.geometry()
        if geometry:
            geometry_paragraph.add_run(f"Tipo: ").bold = True
            geometry_paragraph.add_run(self.get_geometry_type_name(layer.geometryType()))
            geometry_paragraph.add_run("\n")
            
            if layer.geometryType() == QgsWkbTypes.PolygonGeometry:
                # Add area
                area = geometry.area()
                if area < 10000:
                    area_str = f"{area:.2f} m²"
                else:
                    area_str = f"{area/10000:.4f} ha"
                
                geometry_paragraph.add_run(f"Área: ").bold = True
                geometry_paragraph.add_run(area_str)
                geometry_paragraph.add_run("\n")
                
                # Add perimeter
                perimeter = geometry.length()
                if perimeter < 1000:
                    perimeter_str = f"{perimeter:.2f} m"
                else:
                    perimeter_str = f"{perimeter/1000:.4f} km"
                
                geometry_paragraph.add_run(f"Perímetro: ").bold = True
                geometry_paragraph.add_run(perimeter_str)
            
            elif layer.geometryType() == QgsWkbTypes.LineGeometry:
                # Add length
                length = geometry.length()
                if length < 1000:
                    length_str = f"{length:.2f} m"
                else:
                    length_str = f"{length/1000:.4f} km"
                
                geometry_paragraph.add_run(f"Longitud: ").bold = True
                geometry_paragraph.add_run(length_str)
            
            elif layer.geometryType() == QgsWkbTypes.PointGeometry:
                # Add coordinates
                point = geometry.asPoint()
                
                geometry_paragraph.add_run(f"Coordenadas: ").bold = True
                geometry_paragraph.add_run(f"X: {point.x():.6f}, Y: {point.y():.6f}")
    
    def generate_map_image(self):
        """Generate map image of search results."""
        # Create temporary file for map image
        map_image_path = os.path.join(self.temp_dir, "map_image.png")
        
        # Get all features from results
        features_by_layer = {}
        for result in self.results:
            layer = result['layer']
            feature_id = result['feature_id']
            
            if layer not in features_by_layer:
                features_by_layer[layer] = []
            
            features_by_layer[layer].append(feature_id)
        
        # Calculate combined extent
        combined_extent = None
        
        for layer, feature_ids in features_by_layer.items():
            for feature_id in feature_ids:
                # Get feature
                request = QgsFeatureRequest().setFilterFid(feature_id)
                features = list(layer.getFeatures(request))
                
                if features:
                    feature = features[0]
                    geometry = feature.geometry()
                    
                    if geometry:
                        # Get feature extent
                        extent = geometry.boundingBox()
                        
                        # Transform extent to project CRS if needed
                        if layer.crs() != QgsProject.instance().crs():
                            transform = QgsCoordinateTransform(
                                layer.crs(),
                                QgsProject.instance().crs(),
                                QgsProject.instance()
                            )
                            extent = transform.transformBoundingBox(extent)
                        
                        # Combine extents
                        if combined_extent is None:
                            combined_extent = extent
                        else:
                            combined_extent.combineExtentWith(extent)
        
        if not combined_extent:
            return None
        
        # Add buffer for better visibility
        combined_extent.grow(combined_extent.width() * 0.1)
        
        # Create map settings
        map_settings = QgsMapSettings()
        map_settings.setDestinationCrs(QgsProject.instance().crs())
        map_settings.setExtent(combined_extent)
        
        # Set output size
        map_settings.setOutputSize(QSize(800, 600))
        
        # Add all visible layers
        layers = QgsProject.instance().mapLayers().values()
        map_settings.setLayers(layers)
        
        # Create map renderer
        render_job = QgsMapRendererParallelJob(map_settings)
        render_job.start()
        render_job.waitForFinished()
        
        # Get rendered image
        image = render_job.renderedImage()
        
        # Save image
        image.save(map_image_path, "PNG")
        
        return map_image_path
    
    def generate_charts(self, selected_fields):
        """Generate charts for selected fields."""
        chart_paths = {}
        
        # Group results by layer
        results_by_layer = {}
        for result in self.results:
            layer = result['layer']
            if layer not in results_by_layer:
                results_by_layer[layer] = []
            results_by_layer[layer].append(result)
        
        # Generate charts for each field
        for field_name in selected_fields:
            # Check if field exists in any layer
            field_exists = False
            for layer in results_by_layer.keys():
                if field_name in layer.fields().names():
                    field_exists = True
                    break
            
            if not field_exists:
                continue
            
            # Count values across all layers
            value_counts = {}
            
            for layer, layer_results in results_by_layer.items():
                if field_name not in layer.fields().names():
                    continue
                
                # Get unique features
                unique_features = set(r['feature_id'] for r in layer_results)
                
                for feature_id in unique_features:
                    request = QgsFeatureRequest().setFilterFid(feature_id)
                    features = list(layer.getFeatures(request))
                    
                    if features:
                        feature = features[0]
                        value = str(feature[field_name])
                        
                        if value in value_counts:
                            value_counts[value] += 1
                        else:
                            value_counts[value] = 1
            
            # Skip if no values
            if not value_counts:
                continue
            
            # Sort by count
            sorted_items = sorted(value_counts.items(), key=lambda x: x[1], reverse=True)
            
            # Limit to top 10 for readability
            if len(sorted_items) > 10:
                sorted_items = sorted_items[:10]
                other_count = sum(value_counts[k] for k in value_counts if k not in [x[0] for x in sorted_items])
                if other_count > 0:
                    sorted_items.append(("Otros", other_count))
            
            # Extract labels and values
            labels = [item[0] for item in sorted_items]
            values = [item[1] for item in sorted_items]
            
            # Create chart
            plt.figure(figsize=(8, 6))
            
            # Create bar chart
            plt.bar(labels, values)
            plt.xlabel('Valores')
            plt.ylabel('Conteo')
            plt.title(f'Distribución de valores para el campo "{field_name}"')
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            
            # Save chart
            chart_path = os.path.join(self.temp_dir, f"chart_{field_name}.png")
            plt.savefig(chart_path)
            plt.close()
            
            chart_paths[field_name] = chart_path
        
        return chart_paths
    
    def convert_to_pdf(self, docx_path, pdf_path):
        """Convert Word document to PDF."""
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        except ImportError:
            QMessageBox.warning(
                self,
                "Advertencia",
                "La biblioteca docx2pdf no está instalada. "
                "El reporte se ha guardado en formato Word."
            )
            # Copy the Word file to the requested PDF path but with .docx extension
            import shutil
            pdf_base = os.path.splitext(pdf_path)[0]
            shutil.copy2(docx_path, f"{pdf_base}.docx")
    
    def get_geometry_type_name(self, geometry_type):
        """Get human-readable name for geometry type."""
        if geometry_type == QgsWkbTypes.PointGeometry:
            return "Punto"
        elif geometry_type == QgsWkbTypes.LineGeometry:
            return "Línea"
        elif geometry_type == QgsWkbTypes.PolygonGeometry:
            return "Polígono"
        else:
            return "Desconocido"
